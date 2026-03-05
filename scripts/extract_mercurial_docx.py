#!/usr/bin/env python3
"""
Extraction des données du Mercurial des prix (Burkina Faso) depuis un fichier Word (.docx).
Utilise python-docx pour lire les tableaux natifs et exporter en CSV pour import Prisma.
Plus fiable que l'extraction PDF car les tableaux Word ont une structure explicite.
"""

import csv
import re
import sys
from pathlib import Path

from docx import Document


# Format Mercuriale BF : Code XX.X.X.X.X.X.XXX (7 segments)
CODE_PATTERN = re.compile(r"^\d{2}(\.\d{1,2}){5}\.\d{1,3}$")
# Codes courts (catégories) : 03.1.1, 03.1.1.1, etc.
CODE_SHORT_PATTERN = re.compile(r"^\d{2}(\.\d{1,2}){0,5}$")

# Mots-clés pour détecter les en-têtes de colonnes
HEADER_KEYWORDS = {
    "code": ["code", "article"],
    "designation": ["désignation", "designation", "libellé", "libelle", "caractéristiques"],
    "unite": ["unité", "unite", "conditionnement", "condition"],
    "prix_min": ["minimum", "min", "mini"],
    "prix_moyen": ["moyen", "indicatif"],
    "prix_max": ["maximum", "max", "plafond", "prix unitaire plafond"],
}


def clean_price(value: str) -> str | None:
    """
    Nettoie un prix FCFA : espaces (42 500), virgule décimale (42,50), point milliers (42.500 ou 1.500.000).
    """
    if not value or not value.strip():
        return None
    s = value.strip().replace(" ", "").replace("\u00a0", "")
    s = re.sub(r"[^\d.,\-]", "", s)
    if not s or s in ("-", "."):
        return None
    parts_comma = s.split(",")
    parts_dot = s.split(".")
    if "," in s and "." not in s:
        if len(parts_comma) == 2 and len(parts_comma[1]) == 3 and parts_comma[1].isdigit():
            s = "".join(parts_comma)
        else:
            s = s.replace(",", ".")
    elif "." in s and "," not in s:
        if len(parts_dot) >= 2 and all(p.isdigit() for p in parts_dot) and len(parts_dot[-1]) == 3:
            s = "".join(parts_dot)
    elif "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    else:
        s = s.replace(",", ".")
    s = re.sub(r"[^\d.\-]", "", s)
    try:
        float(s)
        return s
    except ValueError:
        return None


def clean_text(value: str) -> str:
    """Nettoie une chaîne : espaces multiples."""
    if not value:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())


def get_column_indices(headers: list[str]) -> dict[str, int]:
    """
    Détecte les indices des colonnes d'après les en-têtes.
    Retourne un dict {code: 0, designation: 1, unite: 2, prix_min: 3, prix_moyen: 4, prix_max: 5}
    """
    indices = {}
    headers_lower = [str(h).lower().strip() for h in headers]

    for col_name, keywords in HEADER_KEYWORDS.items():
        for i, h in enumerate(headers_lower):
            if any(kw in h for kw in keywords):
                indices[col_name] = i
                break
        if col_name not in indices:
            indices[col_name] = -1

    # Fallback : ordre classique Code, Désignation, Unité, Min, Moyen, Max
    if indices["code"] < 0:
        indices["code"] = 0
    if indices["designation"] < 0:
        indices["designation"] = min(1, len(headers) - 1)
    if indices["unite"] < 0:
        indices["unite"] = min(2, len(headers) - 1)
    # Dernières colonnes numériques = Min, Moyen, Max
    if indices["prix_min"] < 0 and indices["prix_moyen"] < 0 and indices["prix_max"] < 0:
        if len(headers) >= 6:
            indices["prix_min"], indices["prix_moyen"], indices["prix_max"] = (
                len(headers) - 3,
                len(headers) - 2,
                len(headers) - 1,
            )
        elif len(headers) >= 4:
            indices["prix_max"] = len(headers) - 1  # Prix plafond seul

    return indices


def is_header_row(cells: list[str]) -> bool:
    """Détecte si une ligne est un en-tête de tableau."""
    if len(cells) < 2:
        return False
    text = " ".join(cells).lower()
    return any(
        kw in text
        for kw in ["code", "désignation", "designation", "libellé", "unité", "conditionnement", "minimum", "maximum"]
    )


def extract_from_table(table, region_id: str, table_index: int) -> list[dict]:
    """
    Extrait les lignes d'un tableau Word.
    Chaque cellule est lue directement : table.rows[i].cells[j].text
    """
    rows = []
    if not table.rows:
        return rows

    # Première ligne = en-têtes
    header_cells = [cell.text for cell in table.rows[0].cells]
    indices = get_column_indices(header_cells)
    start_row = 1

    # Si la première ligne ne ressemble pas à un en-tête, traiter toute la table
    if not is_header_row(header_cells):
        start_row = 0
        # Deviner les indices par position (Code=0, Désignation=1, etc.)
        num_cols = len(table.rows[0].cells)
        indices = {
            "code": 0,
            "designation": min(1, num_cols - 1),
            "unite": min(2, num_cols - 1),
            "prix_min": min(3, num_cols - 1) if num_cols >= 6 else -1,
            "prix_moyen": min(4, num_cols - 1) if num_cols >= 6 else -1,
            "prix_max": num_cols - 1 if num_cols >= 4 else -1,
        }

    for row_idx in range(start_row, len(table.rows)):
        row = table.rows[row_idx]
        cells = [cell.text.strip() for cell in row.cells]

        # Compléter si fusion de cellules (python-docx peut avoir des cellules en moins)
        while len(cells) < max(indices.values()) + 1 and max(indices.values()) >= 0:
            cells.append("")

        def get_cell(idx: int) -> str:
            return cells[idx] if 0 <= idx < len(cells) else ""

        code = clean_text(get_cell(indices["code"]))
        designation = clean_text(get_cell(indices["designation"]))
        unite = clean_text(get_cell(indices["unite"])) or "Unité"

        prix_min = clean_price(get_cell(indices["prix_min"])) if indices["prix_min"] >= 0 else None
        prix_moyen = clean_price(get_cell(indices["prix_moyen"])) if indices["prix_moyen"] >= 0 else None
        prix_max = clean_price(get_cell(indices["prix_max"])) if indices["prix_max"] >= 0 else None

        # Si une seule colonne prix : c'est le plafond
        if not prix_min and not prix_moyen and not prix_max:
            continue
        if not prix_max and (prix_min or prix_moyen):
            prix_max = prix_moyen or prix_min

        # Ignorer les lignes sans code ni désignation valide
        if not code and not designation:
            continue
        # Ignorer les lignes d'en-tête répétées
        if is_header_row(cells):
            continue
        # Ignorer les lignes sans prix (sous-titres de section)
        if not prix_min and not prix_moyen and not prix_max:
            continue

        def to_str(v) -> str:
            if not v:
                return ""
            try:
                f = float(v)
                return str(int(f)) if f == int(f) else str(f)
            except (ValueError, TypeError):
                return ""

        rows.append({
            "code": code or designation[:50] if designation else "",
            "designation": designation or code,
            "unite": unite,
            "prix_min": to_str(prix_min),
            "prix_moyen": to_str(prix_moyen),
            "prix_max": to_str(prix_max),
            "prix_unitaire_plafond": to_str(prix_max) or to_str(prix_moyen) or to_str(prix_min),
            "region_id": region_id,
            "page": table_index + 1,
        })

    return rows


def process_docx(docx_path: Path, output_path: Path, region_id: str = "centre") -> int:
    """
    Parcourt le document Word, extrait les tableaux et exporte en CSV.
    Retourne le nombre de lignes extraites.
    """
    doc = Document(docx_path)
    all_rows = []

    for table_idx, table in enumerate(doc.tables):
        table_rows = extract_from_table(table, region_id, table_idx)
        all_rows.extend(table_rows)
        if (table_idx + 1) % 10 == 0 and table_idx > 0:
            print(f"  Tableau {table_idx + 1}/{len(doc.tables)} - {len(all_rows)} lignes")

    # Déduplication (code + conditionnement)
    seen = set()
    unique_rows = []
    for row in all_rows:
        key = (row["code"], row["unite"])
        if key in seen:
            continue
        seen.add(key)
        # Filtrer les codes invalides (lignes parasites)
        if row["code"] and not re.match(r"^[\d.]+$", row["code"]) and len(row["code"]) < 3:
            continue
        unique_rows.append(row)

    # Export CSV
    csv_columns = [
        "code", "designation", "unite",
        "prix_min", "prix_moyen", "prix_max", "prix_unitaire_plafond",
        "region_id", "page"
    ]
    with open(output_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=csv_columns, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(unique_rows)

    return len(unique_rows)


def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_mercurial_docx.py <chemin_mercurial.docx> [region_id] [output.csv]")
        print("Exemple: python extract_mercurial_docx.py Mercuriale_2026_OUAGADOUGOU.docx ouagadougou scripts/mercurial_ouagadougou.csv")
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Erreur: fichier introuvable: {docx_path}")
        sys.exit(1)
    if docx_path.suffix.lower() not in (".docx",):
        print("Erreur: le fichier doit être au format .docx")
        sys.exit(1)

    region_id = sys.argv[2] if len(sys.argv) > 2 else "centre"
    output_path = Path(sys.argv[3]) if len(sys.argv) > 3 else docx_path.with_suffix(".csv")

    print(f"Extraction du Mercurial Word: {docx_path}")
    print(f"Région: {region_id}")
    print(f"Sortie: {output_path}")
    print("-" * 50)

    count = process_docx(docx_path, output_path, region_id)

    print("-" * 50)
    print(f"Terminé: {count} lignes exportées vers {output_path}")


if __name__ == "__main__":
    main()
