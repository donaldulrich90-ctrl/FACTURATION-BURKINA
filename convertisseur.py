import docx
import json
import os
import pandas as pd
import re
import sys
from tqdm import tqdm

# Mots-clés pour détecter les lignes d'en-tête
HEADER_KEYWORDS = ["code", "désignation", "designation", "conditionnement", "minimum", "moyen", "maximum"]

# Plafond prix FCFA (limite INT SQLite/Prisma) - évite valeurs aberrantes
PRIX_MAX_VALID = 999_999_999

def clean_price(price_str):
    """
    Parse le prix FCFA exactement comme dans le Word.
    - Espaces milliers : "8 000" -> 8000, "89 000" -> 89000
    - Point milliers (européen) : "42.500" -> 42500, "1.500.000" -> 1500000
    - Virgule décimale : "42,50" -> 42.5
    """
    if not price_str: return None
    # Supprimer tous les espaces (normaux, insécables, fins, etc.) pour "8 000" -> "8000"
    s = str(price_str).strip()
    for sp in (" ", "\xa0", "\u00a0", "\u2009", "\u202f", "\u2003", "\t"):
        s = s.replace(sp, "")
    s = re.sub(r"[^\d.,\-]", "", s)
    if not s or s in ("-", "."): return None
    # Virgule seule = séparateur décimal ou milliers
    if "," in s and "." not in s:
        parts = s.split(",")
        if len(parts) == 2 and len(parts[1]) == 3 and parts[1].isdigit():
            s = "".join(parts)  # 42,500 -> 42500
        else:
            s = s.replace(",", ".")
    # Point seul = milliers européen (1.500.000) ou décimal
    elif "." in s and "," not in s:
        parts = s.split(".")
        if len(parts) >= 2 and all(p.isdigit() for p in parts) and len(parts[-1]) == 3:
            s = "".join(parts)  # 1.500.000 -> 1500000
    elif "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    try:
        val = float(s.replace(",", ".")) if s else None
        if val is not None and (val < 0 or val > PRIX_MAX_VALID): return None
        return int(val) if val is not None and val == int(val) else val
    except (ValueError, TypeError):
        return None

def is_header_row(cells):
    """Détecte si la ligne est un en-tête."""
    text = " ".join(str(c).lower() for c in cells[:6])
    return any(kw in text for kw in HEADER_KEYWORDS)

# Mots à exclure de la détection de région (en-têtes de colonnes)
EXCLUS_REGION = {"minimum", "moyen", "maximum", "conditionnement", "unité", "unite", "code", "désignation", "designation", "caractéristiques"}

def detect_region_and_indices(table):
    """
    Détecte la région (ex: OUAGADOUGOU) et la structure des colonnes.
    Structure attendue: Code | Désignation/Caractéristiques | Conditionnement | Min | Moyen | Max
    La région doit être un nom de lieu (lettres + tirets), pas un en-tête de colonne.
    """
    region = "OUAGADOUGOU"  # défaut
    for row in table.rows[:4]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 6:
            for c in cells[3:6]:
                # Région = lettres+tirets uniquement, pas un en-tête connu
                if c and len(c) > 3 and re.match(r'^[A-Za-zÀ-ÿ\-]+$', c):
                    if c.lower() not in EXCLUS_REGION:
                        region = c
                        break
            low = " ".join(c.lower() for c in cells)
            if "minimum" in low and "moyen" in low and "maximum" in low:
                break
    return region

def get_region_from_path(word_path):
    """Détecte la région depuis le nom du fichier ou le premier tableau."""
    doc = docx.Document(word_path)
    fname = os.path.basename(str(word_path)).upper()
    region = "OUAGADOUGOU"
    for r in ["OUAGADOUGOU", "BOBO-DIOULASSO", "OUAHIGOUYA", "KAYA", "FADA", "TENKODOGO", "DEDOUGOU", "BANFORA", "KOUDOUGOU", "CENTRE", "HAUTS-BASSINS", "CASCADES"]:
        if r.replace("-", "_") in fname.replace("-", "_") or r in fname:
            return r
    if doc.tables and doc.tables[0].rows:
        return detect_region_and_indices(doc.tables[0])
    return region


def word_to_lines(word_path, silent=False):
    """
    Extrait les lignes mercuriale depuis un fichier Word.
    Retourne une liste de dicts au format plateforme: code, designation, conditionnement, prix_min, prix_moyen, prix_max, type, categorie.
    """
    doc = docx.Document(word_path)
    data = []
    current_categorie = "Divers"
    tables_iter = tqdm(doc.tables, disable=silent, desc="Extraction") if not silent else doc.tables
    for table in tables_iter:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            while len(cells) < 6:
                cells.append("")

            if len(cells) < 6 or is_header_row(cells):
                continue

            code = cells[0]
            designation = cells[1]
            conditionnement = cells[2]
            prix_min = clean_price(cells[3]) if len(cells) > 3 else None
            prix_moyen = clean_price(cells[4]) if len(cells) > 4 else None
            prix_max = clean_price(cells[5]) if len(cells) > 5 else None

            if not code and not designation:
                continue

            has_prix = prix_min is not None or prix_moyen is not None or prix_max is not None
            def cap(v):
                if v is None: return 0
                return min(max(0, int(v)), PRIX_MAX_VALID)
            p_min = cap(prix_min) if prix_min is not None else (cap(prix_moyen) if prix_moyen else (cap(prix_max) if prix_max else 0))
            p_moy = cap(prix_moyen) if prix_moyen is not None else (cap(prix_min) if prix_min else (cap(prix_max) if prix_max else 0))
            p_max = cap(prix_max) if prix_max is not None else (cap(prix_moyen) if prix_moyen else (cap(prix_min) if prix_min else 0))

            if has_prix:
                data.append({
                    "type": "article",
                    "code": code,
                    "designation": designation,
                    "conditionnement": conditionnement or "Unité",
                    "prix_min": p_min,
                    "prix_moyen": p_moy,
                    "prix_max": p_max,
                    "categorie": current_categorie,
                })
            else:
                current_categorie = (designation or code or "Divers")[:200]
                data.append({
                    "type": "category",
                    "code": code or f"CAT-{len(data)}",
                    "designation": designation or code,
                    "categorie": current_categorie,
                })

    # Dédupliquer par (code, conditionnement)
    seen = set()
    unique = []
    for row in data:
        key = (row.get("code", ""), row.get("conditionnement", ""))
        if key in seen and row.get("type") == "article":
            continue
        seen.add(key)
        unique.append(row)
    return unique


def word_to_csv(word_path, csv_path):
    print(f"Ouverture du fichier : {word_path}")
    lines = word_to_lines(word_path, silent=False)
    region = get_region_from_path(word_path)
    articles = [l for l in lines if l.get("type") == "article"]
    df = pd.DataFrame([
        {"Code": l["code"], "Designation_Caracteristiques": l["designation"], "Conditionnement": l.get("conditionnement", ""),
         f"Prix_{region}_Minimum": l.get("prix_min"), f"Prix_{region}_Moyen": l.get("prix_moyen"), f"Prix_{region}_Maximum": l.get("prix_max")}
        for l in articles
    ])
    df = df.drop_duplicates()
    df.to_csv(csv_path, index=False, sep=';', encoding='utf-8-sig')
    print(f"\nTerminé ! {len(df)} articles exportés dans {csv_path}")

# Lancement
if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    flags = [a for a in sys.argv[1:] if a.startswith("--")]
    json_mode = "--json" in flags

    if len(args) >= 1:
        word_path = args[0]
        csv_path = args[1] if len(args) >= 2 else "mercurial_pret_a_nettoyer.csv"
    else:
        word_path = r"C:\Users\user\Downloads\Mercuriale_2026_VersionFinale\Mercuriale_2026_VersionFinale\VERSION WORD\Mercuriale_2026_OUAGADOUGOU.docx"
        csv_path = "mercurial_pret_a_nettoyer.csv"

    if json_mode:
        # Sortie JSON pour intégration plateforme (serveur Node.js)
        lines = word_to_lines(word_path, silent=True)
        try:
            out = json.dumps(lines, ensure_ascii=False)
            sys.stdout.buffer.write(out.encode('utf-8'))
        except (UnicodeEncodeError, AttributeError):
            sys.stdout.write(json.dumps(lines, ensure_ascii=True))
    else:
        word_to_csv(word_path, csv_path)